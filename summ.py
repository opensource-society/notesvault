"""
NotesVault Summarizer
-----------------------
A simple tool to summarize notes from .txt, .pdf, or .docx files.

✨ Features:
- Choose which file to summarize (txt, pdf, docx)
- Choose summary length (Short = 8, Medium = 12, Detailed = 20 sentences)
- View summary in the terminal
- Save summary with a custom file name
- Save in .txt, .docx, or .pdf format
- Download/share ready
"""

import os
from docx import Document
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# ------------------ File Reading Functions ------------------ #
def read_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def read_docx(file_path):
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

# ------------------ Summarizer (Dummy Example) ------------------ #
def generate_summary(text, length_choice):
    sentences = text.split(".")
    sentences = [s.strip() for s in sentences if s.strip()]  

    if length_choice == "1":
        return ". ".join(sentences[:8]) + "."
    elif length_choice == "2":
        return ". ".join(sentences[:12]) + "."
    elif length_choice == "3":
        return ". ".join(sentences[:20]) + "."
    else:
        return ". ".join(sentences[:12]) + "."

# ------------------ Save Functions ------------------ #
def save_as_txt(summary, filename):
    with open(filename + ".txt", "w", encoding="utf-8") as f:
        f.write(summary)

def save_as_docx(summary, filename):
    doc = Document()
    doc.add_paragraph(summary)
    doc.save(filename + ".docx")

def save_as_pdf(summary, filename):
    c = canvas.Canvas(filename + ".pdf", pagesize=letter)
    text_obj = c.beginText(50, 750)
    text_obj.setFont("Helvetica", 12)

    for line in summary.split(". "):
        text_obj.textLine(line.strip() + ".")
    c.drawText(text_obj)
    c.save()

# ------------------ Main Program ------------------ #
def main():
    print("📝 NotesVault - Summarizer")
    print("----------------------------------------")

    # Step 1: File input
    file_path = input("Enter your notes file path (PDF, DOCX, TXT): ").strip()

    if not os.path.exists(file_path):
        print("❌ File not found. Please check the path and try again.")
        return

    if file_path.endswith(".txt"):
        text = read_txt(file_path)
    elif file_path.endswith(".docx"):
        text = read_docx(file_path)
    elif file_path.endswith(".pdf"):
        text = read_pdf(file_path)
    else:
        print("❌ Unsupported file format. Use TXT, DOCX, or PDF.")
        return

    # Step 2: Choose summary length
    print("\n📌 Choose summary length:")
    print("1. Short (8 sentences)")
    print("2. Medium (12 sentences)")
    print("3. Detailed (20 sentences)")
    length_choice = input("Enter 1 / 2 / 3: ").strip()

    summary = generate_summary(text, length_choice)

    # Step 3: Show summary
    print("\n✅ Generated Summary:\n")
    print(summary)

    # Step 4: Save summary
    save_choice = input("\n💾 Do you want to save the summary? (y/n): ").lower()
    if save_choice == "y":
        filename = input("Enter a file name for your summary (without extension): ").strip()
        print("📂 Choose format to save in:")
        print("1. TXT")
        print("2. DOCX")
        print("3. PDF")
        format_choice = input("Enter 1 / 2 / 3: ").strip()

        if format_choice == "1":
            save_as_txt(summary, filename)
            print(f"✅ Summary saved as {filename}.txt")
        elif format_choice == "2":
            save_as_docx(summary, filename)
            print(f"✅ Summary saved as {filename}.docx")
        elif format_choice == "3":
            save_as_pdf(summary, filename)
            print(f"✅ Summary saved as {filename}.pdf")
        else:
            print("❌ Invalid choice. Summary not saved.")

    print("\n✨ Done! Thank you for using NotesVault Summarizer.")

if __name__ == "__main__":
    main()
